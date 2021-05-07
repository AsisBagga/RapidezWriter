from django.shortcuts import render, redirect, HttpResponse
from django.conf import settings
from django.core.files.storage import FileSystemStorage
from .forms import ResumeForm
from .models import Resume
from django.contrib import messages

#Email stuff
from django.template.loader import render_to_string
from django.utils.html import strip_tags
from django.core.mail import EmailMultiAlternatives, EmailMessage

# Celery
# from .tasks import sleepy, send_email

# Emails
from django.core.mail import send_mail

# Resume analysis
import docx

# Resume parser
from pyresparser import ResumeParser

import json


def resume_analysis(request):
    return render(request, 'resume_analysis.html')

def thankyou(request):
    return render(request, 'thankyou.html')

def resume_upload(request):
    if request.method == 'POST':
        print("Inside resume_upload func ")
        print("Received POST call")
        form = ResumeForm(request.POST, request.FILES)
        if form.is_valid():
            print("forms valid")
            userResume = form.save()
            userName = userResume.name
            userPhone = userResume.phone
            userEmail = userResume.email
            
            userResumePath = userResume.resume.path
            resume_parser(userResumePath, userName, userEmail, userPhone)
            
            return redirect('resumeAnalysis:thankyou')
        else:
            print("Invalid form")
            print(form.errors)
    else:
            print("POST method Else block")
            form = ResumeForm()
    return render(request, 'resume_upload.html', {'form': form})


def resume_parser(resumePath, userName, userEmail, userPhone):
    doc = docx.Document(resumePath)

    imagePresentInResume = False
    numberOfPages = 0
    numberOfBullets = 0
    sectionCount = 0
    sections = ['skills', 'education', 'experience', 'objective', 'certification']
    secntionsPresent = []

    for paragraph in doc.paragraphs:
        # If you find an image
        if 'Graphic' in paragraph._p.xml:
            imagePresentInResume = True
        else:
            imagePresentInResume = False
        
        if 'w:lastRenderedPageBreak' in paragraph._p.xml:
            numberOfPages += 1
        
        if 'ListBullet' in paragraph._p.xml:
            numberOfBullets +=1

        for section in sections:
            if section in paragraph._p.xml:
                sectionCount +=1
                secntionsPresent.append(section)

    print("numberOfPages:")
    numberOfPages = numberOfPages + 1
    print(numberOfPages)
    if imagePresentInResume == True:
        print("Images or graphic is present in the resume")
    else:
        print("Images or graphic is not present in the resume")

    if numberOfBullets:
        print("Bullets is present in the resume")
        print("numberOfBullets:")
        print(numberOfBullets)
    else:
        print("Bullets is not present in the resume")

    if sectionCount < 2:
        print("Relevant sections are missing")
    elif sectionCount < 4:
        print("Relevant sections are present, maybe can include more optional ones")
    else:
        print("Sections: ")
        print(sectionCount)
        print("Relevant sections included")

    # Email 
    email_from = settings.EMAIL_FROM
    email_admin = settings.EMAIL_ADMIN
    content = 'Resume Analysis'
    

    html_content = render_to_string('mail_template.html', 
                                    {
                                    'name' : userName,
                                    'email': userEmail,
                                    'phone': userPhone,
                                    'sections' : sectionCount,
                                    'imagePresentInResume' : imagePresentInResume,
                                    'numberOfPages' : numberOfPages,
                                    'numberOfBullets' : numberOfBullets,
                                   })
    text_content = strip_tags(html_content)

    email_to_user = EmailMultiAlternatives(
        'Resume Analysis',
        text_content,
        email_temp_from,
        [email_to], #user_email
    )
    email_to_user.attach_alternative(html_content, "text/html")
    email_to_user.send()
    
    resume_summary = "\n Name: "+userName+"\n Email: "+userEmail+"\n Phone Number: "+str(userPhone)+"\n sectionCount: "+str(sectionCount)+"\n imagePresentInResume: "+str(imagePresentInResume)+"\n numberOfPages: "+str(numberOfPages)+"\n numberOfBullets: "+str(numberOfBullets)

    email_to_admin = EmailMessage(
        'Customer Resume Analysis Summary',
        resume_summary,
        email_from,
        [email_from], #email_admin
        )
    
    print(resumePath)
    email_to_admin.attach_file(resumePath)
    email_to_admin.send()
        
    print("EMail sent to admin", email_to_admin)

    return


def test(request):
    print("Inside test function")
    # send_email()

    # doc = docx.Document(r"C:\Stuff\WDF\rapidez\RapidezWriter-1\Samples\Bhattacharya.docx")
    # all_paras = doc.paragraphs
    # print(len(all_paras))
    # return redirect('resumeAnalysis:thankyou')

    # data = ResumeParser(r"C:\Users\bharath878\Downloads\b-suryanarayanan-cv-2020.pdf").get_extracted_data()
    print("\n")
    pretty_data = json.dumps(data, sort_keys=True, indent=4)
    print(pretty_data)
    return render(request, 'thankyou.html', {'pretty_data': pretty_data})
